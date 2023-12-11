import requests
import json
import math
import argparse
import configparser
import warnings
import urllib3

from tqdm import tqdm
from docx import Document

# 抑制 InsecureRequestWarning
warnings.simplefilter('ignore', urllib3.exceptions.InsecureRequestWarning)

def pa_shuoshuo(config):
    url = "https://user.qzone.qq.com/proxy/domain/taotao.qq.com/cgi-bin/emotion_cgi_msglist_v6"
    headers = {
        f'Reference': 'https://user.qzone.qq.com/{config["shuoshuo"]["qq"]}',
        'Accept': '*/*',
        'Accept-Encoding': 'gzip, deflate, br',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'Content-Type': 'application/json; charset=utf-8',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    }
    headers['Cookie'] = config["shuoshuo"]["cookie"]
    params = {
        'hostUin': config["shuoshuo"]["qq"],
        'uin': config["shuoshuo"]["qq"],
        'blogType': 0,
        'reqInfo': 1,
        'pos': 0,
        'num': 10,
        'sortType': 0,
        'inCharset': 'utf-8',
        'outCharset': 'utf-8',
        'g_tk': config["shuoshuo"]["g_tk"]
    }

    # 循环获取所有说说
    doc = Document()
    doc.add_heading('我的说说', 0)

    num_shuoshuo = config.getint("shuoshuo", "number")
    num_loops = math.ceil(num_shuoshuo / 10)
    for i in tqdm(range(num_loops), desc="爬取说说"):
        try:
            params['pos'] = i * 10
            if (i + 1) * 10 > num_shuoshuo:
                params['num'] = num_shuoshuo - i * 10

            res = requests.get(url, params=params, headers=headers, verify=False)
            response = res.text.replace("_Callback(", "").replace(");", "")

            json_data = json.loads(response)
            for msg_data in json_data['msglist']:
                content = msg_data['content']
                create_time = msg_data['createTime']
                city_name = msg_data['lbs']['name']

                doc.add_heading(create_time + " " + city_name, level=1)
                doc.add_paragraph(content)

            doc.save('shuoshuo.docx')
        except Exception as e:
            continue

    doc.save('shuoshuo.docx')


if __name__ == '__main__':
    # 命令行参数
    parser = argparse.ArgumentParser()
    parser.add_argument('-c', '--config', type=str, default="config.ini", help='配置文件')
    args = parser.parse_args()

    # 读取配置文件
    config = configparser.ConfigParser()
    config.read(args.config)

    # 爬取说说
    pa_shuoshuo(config)